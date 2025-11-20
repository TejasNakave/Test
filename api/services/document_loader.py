"""
Document loader for extracting text and images from DOCX and PDF files in the data folder.
Enhanced with multimodal capabilities for Trade Assistant.
"""
import os
import json
import zipfile
from typing import List, Dict
from docx import Document
from langchain.schema import Document as LangChainDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

# PDF processing imports
try:
    import PyPDF2
    import fitz  # PyMuPDF
    from pdfminer.high_level import extract_text as pdfminer_extract
    PDF_SUPPORT = True
    logger.info("✅ PDF processing libraries loaded successfully")
except ImportError as e:
    PDF_SUPPORT = False
    logger.warning(f"⚠️ PDF processing libraries not available: {e}")
    logger.info("📦 Install with: pip install PyPDF2 PyMuPDF pdfminer.six")

logger = logging.getLogger(__name__)

class DocumentLoader:
    def __init__(self, data_folder: str = "data"):
        self.data_folder = data_folder
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        self.extracted_images_folder = "extracted_images"
        self.image_metadata_file = "image_metadata.json"
        
        # Create folders if they don't exist
        os.makedirs(self.extracted_images_folder, exist_ok=True)
    
    def extract_images_from_docx(self, file_path: str, doc_name: str) -> List[Dict]:
        """Extract images from a DOCX file and save them locally."""
        images_info = []
        
        try:
            # DOCX files are essentially ZIP archives
            with zipfile.ZipFile(file_path, 'r') as zip_file:
                # Find all image files in the media folder
                media_files = [f for f in zip_file.namelist() if f.startswith('word/media/')]
                
                for i, media_file in enumerate(media_files):
                    try:
                        # Extract file extension
                        file_extension = os.path.splitext(media_file)[1].lower()
                        
                        # Only process common image formats
                        if file_extension in ['.png', '.jpg', '.jpeg', '.gif', '.bmp']:
                            # Create a unique filename
                            image_filename = f"{doc_name}_image{i+1}{file_extension}"
                            image_path = os.path.join(self.extracted_images_folder, image_filename)
                            
                            # Extract and save the image
                            with zip_file.open(media_file) as source, open(image_path, 'wb') as target:
                                target.write(source.read())
                            
                            # Store image metadata
                            image_info = {
                                'image_filename': image_filename,
                                'image_path': image_path,
                                'original_path': media_file,
                                'file_extension': file_extension,
                                'file_size': os.path.getsize(image_path)
                            }
                            
                            images_info.append(image_info)
                            logger.info(f"Extracted image: {image_filename} from {doc_name}")
                    
                    except Exception as e:
                        logger.error(f"Error extracting image {media_file} from {doc_name}: {e}")
                        continue
        
        except Exception as e:
            logger.error(f"Error extracting images from {file_path}: {e}")
        
        return images_info
    
    def save_image_metadata(self, all_images_metadata: Dict):
        """Save image metadata to JSON file."""
        try:
            with open(self.image_metadata_file, 'w', encoding='utf-8') as f:
                json.dump(all_images_metadata, f, indent=2, ensure_ascii=False)
            logger.info(f"Image metadata saved to {self.image_metadata_file}")
        except Exception as e:
            logger.error(f"Error saving image metadata: {e}")
    
    def load_image_metadata(self) -> Dict:
        """Load existing image metadata."""
        try:
            if os.path.exists(self.image_metadata_file):
                with open(self.image_metadata_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
        except Exception as e:
            logger.error(f"Error loading image metadata: {e}")
        return {}
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text content from a DOCX file."""
        try:
            doc = Document(file_path)
            text = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text.strip())
            
            return "\n".join(text)
        
        except Exception as e:
            logger.error(f"Error reading DOCX file {file_path}: {str(e)}")
            return ""
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text content from a PDF file using multiple methods for better coverage."""
        if not PDF_SUPPORT:
            logger.error("PDF support not available. Install required packages: pip install PyPDF2 PyMuPDF pdfminer.six")
            return ""
        
        text_content = ""
        
        # Method 1: Try PyMuPDF (fitz) - best for complex PDFs
        try:
            logger.info(f"📄 Extracting PDF text with PyMuPDF: {os.path.basename(file_path)}")
            doc = fitz.open(file_path)
            text_parts = []
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                page_text = page.get_text()
                if page_text.strip():
                    text_parts.append(f"[Page {page_num + 1}]\n{page_text.strip()}")
            
            doc.close()
            
            if text_parts:
                text_content = "\n\n".join(text_parts)
                logger.info(f"✅ PyMuPDF extracted {len(text_parts)} pages from {os.path.basename(file_path)}")
                return text_content
            else:
                logger.warning(f"⚠️ PyMuPDF extracted no text from {os.path.basename(file_path)}")
                
        except Exception as e:
            logger.warning(f"⚠️ PyMuPDF failed for {os.path.basename(file_path)}: {str(e)}")
        
        # Method 2: Try PyPDF2 - good for standard PDFs
        try:
            logger.info(f"📄 Extracting PDF text with PyPDF2: {os.path.basename(file_path)}")
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text_parts = []
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_parts.append(f"[Page {page_num + 1}]\n{page_text.strip()}")
                
                if text_parts:
                    text_content = "\n\n".join(text_parts)
                    logger.info(f"✅ PyPDF2 extracted {len(text_parts)} pages from {os.path.basename(file_path)}")
                    return text_content
                else:
                    logger.warning(f"⚠️ PyPDF2 extracted no text from {os.path.basename(file_path)}")
                    
        except Exception as e:
            logger.warning(f"⚠️ PyPDF2 failed for {os.path.basename(file_path)}: {str(e)}")
        
        # Method 3: Try PDFMiner - good for complex layouts
        try:
            logger.info(f"📄 Extracting PDF text with PDFMiner: {os.path.basename(file_path)}")
            text_content = pdfminer_extract(file_path)
            
            if text_content and text_content.strip():
                logger.info(f"✅ PDFMiner extracted text from {os.path.basename(file_path)}")
                return text_content.strip()
            else:
                logger.warning(f"⚠️ PDFMiner extracted no text from {os.path.basename(file_path)}")
                
        except Exception as e:
            logger.warning(f"⚠️ PDFMiner failed for {os.path.basename(file_path)}: {str(e)}")
        
        # If all methods fail
        logger.error(f"❌ All PDF extraction methods failed for {os.path.basename(file_path)}")
        return ""
    
    def extract_images_from_pdf(self, file_path: str, doc_name: str) -> List[Dict]:
        """Extract images from a PDF file and save them locally."""
        if not PDF_SUPPORT:
            logger.warning("PDF image extraction not available without PyMuPDF")
            return []
        
        images_info = []
        
        try:
            logger.info(f"🖼️ Extracting images from PDF: {doc_name}")
            doc = fitz.open(file_path)
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                image_list = page.get_images(full=True)
                
                for img_index, img in enumerate(image_list):
                    try:
                        # Get image data
                        xref = img[0]
                        pix = fitz.Pixmap(doc, xref)
                        
                        # Skip images that are too small or masks
                        if pix.width < 50 or pix.height < 50:
                            pix = None
                            continue
                        
                        # Determine image format
                        if pix.n - pix.alpha < 4:  # GRAY or RGB
                            image_format = "png"
                        else:  # CMYK
                            pix1 = fitz.Pixmap(fitz.csRGB, pix)
                            pix = pix1
                            image_format = "png"
                        
                        # Create filename
                        image_filename = f"{doc_name}_page{page_num+1}_img{img_index+1}.{image_format}"
                        image_path = os.path.join(self.extracted_images_folder, image_filename)
                        
                        # Save image
                        pix.save(image_path)
                        
                        # Store image info
                        image_info = {
                            "filename": image_filename,
                            "source_document": doc_name,
                            "page_number": page_num + 1,
                            "image_index": img_index + 1,
                            "width": pix.width,
                            "height": pix.height,
                            "format": image_format,
                            "file_path": image_path,
                            "extracted_at": str(os.path.getctime(image_path)) if os.path.exists(image_path) else None
                        }
                        
                        images_info.append(image_info)
                        logger.info(f"✅ Extracted image: {image_filename}")
                        
                        pix = None  # Free memory
                        
                    except Exception as e:
                        logger.error(f"Error extracting image {img_index} from page {page_num+1}: {e}")
                        continue
            
            doc.close()
            logger.info(f"🎉 Extracted {len(images_info)} images from {doc_name}")
            
        except Exception as e:
            logger.error(f"Error extracting images from PDF {file_path}: {e}")
        
        return images_info
    
    def load_documents(self) -> List[LangChainDocument]:
        """Load all DOCX and PDF documents from the data folder and extract images."""
        documents = []
        all_images_metadata = self.load_image_metadata()  # Load existing metadata
        
        if not os.path.exists(self.data_folder):
            logger.warning(f"Data folder '{self.data_folder}' not found")
            return documents
        
        # Get both DOCX and PDF files
        supported_files = []
        docx_files = [f for f in os.listdir(self.data_folder) if f.endswith('.docx')]
        pdf_files = [f for f in os.listdir(self.data_folder) if f.endswith('.pdf')] if PDF_SUPPORT else []
        
        supported_files.extend([(f, 'docx') for f in docx_files])
        supported_files.extend([(f, 'pdf') for f in pdf_files])
        
        if not supported_files:
            file_types = "DOCX and PDF" if PDF_SUPPORT else "DOCX"
            logger.warning(f"No {file_types} files found in '{self.data_folder}'")
            return documents
        
        logger.info(f"Found {len(docx_files)} DOCX and {len(pdf_files) if PDF_SUPPORT else 0} PDF files to process")
        
        for file_name, file_type in supported_files:
            file_path = os.path.join(self.data_folder, file_name)
            logger.info(f"Processing {file_type.upper()} file: {file_name}")
            
            # Extract text content based on file type
            if file_type == 'docx':
                text_content = self.extract_text_from_docx(file_path)
            elif file_type == 'pdf':
                text_content = self.extract_text_from_pdf(file_path)
            else:
                logger.warning(f"Unsupported file type for {file_name}")
                continue
            
            # Extract images (only if not already processed)
            doc_name_clean = os.path.splitext(file_name)[0]
            if doc_name_clean not in all_images_metadata:
                logger.info(f"Extracting images from: {file_name}")
                
                if file_type == 'docx':
                    images_info = self.extract_images_from_docx(file_path, doc_name_clean)
                elif file_type == 'pdf':
                    images_info = self.extract_images_from_pdf(file_path, doc_name_clean)
                else:
                    images_info = []
                
                if images_info:
                    all_images_metadata[doc_name_clean] = images_info
                    logger.info(f"Extracted {len(images_info)} images from {file_name}")
                else:
                    all_images_metadata[doc_name_clean] = []  # Mark as processed even if no images
            else:
                logger.info(f"Images already extracted from: {file_name}")
            
            if text_content:
                # Split the document into chunks
                chunks = self.text_splitter.split_text(text_content)
                
                for i, chunk in enumerate(chunks):
                    doc = LangChainDocument(
                        page_content=chunk,
                        metadata={
                            "source": file_name,
                            "chunk_id": i,
                            "file_path": file_path,
                            "file_type": file_type,
                            "total_chunks": len(chunks),
                            "has_images": len(all_images_metadata.get(doc_name_clean, [])) > 0,
                            "image_count": len(all_images_metadata.get(doc_name_clean, []))
                        }
                    )
                    documents.append(doc)
                
                logger.info(f"Created {len(chunks)} chunks from {file_name}")
            else:
                logger.warning(f"No text content extracted from {file_name}")
        
        # Save updated image metadata
        self.save_image_metadata(all_images_metadata)
        
        # Log summary
        total_images = sum(len(images) for images in all_images_metadata.values())
        docs_with_images = sum(1 for images in all_images_metadata.values() if len(images) > 0)
        
        logger.info(f"Total documents loaded: {len(documents)}")
        logger.info(f"Total images extracted: {total_images} from {docs_with_images} documents")
        
        return documents
    
    def get_file_info(self) -> Dict[str, int]:
        """Get information about files in the data folder."""
        info = {
            "total_files": 0,
            "docx_files": 0,
            "pdf_files": 0,
            "supported_files": 0,
            "other_files": 0
        }
        
        if not os.path.exists(self.data_folder):
            return info
        
        files = os.listdir(self.data_folder)
        info["total_files"] = len(files)
        info["docx_files"] = len([f for f in files if f.endswith('.docx')])
        info["pdf_files"] = len([f for f in files if f.endswith('.pdf')])
        info["supported_files"] = info["docx_files"] + (info["pdf_files"] if PDF_SUPPORT else 0)
        info["other_files"] = info["total_files"] - info["supported_files"]
        
        return info
    
    def load_single_document(self, file_path: str) -> List[LangChainDocument]:
        """Load a single DOCX or PDF document."""
        documents = []
        
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return documents
        
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension == '.docx':
            text_content = self.extract_text_from_docx(file_path)
        elif file_extension == '.pdf' and PDF_SUPPORT:
            text_content = self.extract_text_from_pdf(file_path)
        else:
            supported_types = "DOCX and PDF" if PDF_SUPPORT else "DOCX"
            logger.error(f"File type not supported. Supported types: {supported_types}. File: {file_path}")
            return documents
        
        if text_content:
            chunks = self.text_splitter.split_text(text_content)
            
            for i, chunk in enumerate(chunks):
                doc = LangChainDocument(
                    page_content=chunk,
                    metadata={
                        "source": os.path.basename(file_path),
                        "chunk_id": i,
                        "file_path": file_path,
                        "file_type": file_extension,
                        "total_chunks": len(chunks)
                    }
                )
                documents.append(doc)
            
            logger.info(f"Loaded single document: {len(documents)} chunks from {os.path.basename(file_path)}")
        else:
            logger.warning(f"No text content extracted from {file_path}")
        
        return documents
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats."""
        formats = ['.docx']
        if PDF_SUPPORT:
            formats.append('.pdf')
        return formats
    
    def is_pdf_support_available(self) -> bool:
        """Check if PDF support is available."""
        return PDF_SUPPORT