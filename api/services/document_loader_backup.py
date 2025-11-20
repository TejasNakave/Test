"""
Document loader for extracting text from DOCX files in the data folder.
Integrated with ChatGPT-like Proactive Chatbot backend.
"""
import os
from typing import List, Dict
from docx import Document
from langchain.schema import Document as LangChainDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
import logging

logger = logging.getLogger(__name__)

class DocumentLoader:
    def __init__(self, data_folder: str = "data"):
        self.data_folder = data_folder
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text content from a DOCX file."""
        try:
            doc = Document(file_path)
            text = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text.strip())
            
            return "\n".join(text)
        
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return ""
    
    def load_documents(self) -> List[LangChainDocument]:
        """Load all DOCX documents from the data folder."""
        documents = []
        
        if not os.path.exists(self.data_folder):
            logger.warning(f"Data folder '{self.data_folder}' not found")
            return documents
        
        docx_files = [f for f in os.listdir(self.data_folder) if f.endswith('.docx')]
        
        if not docx_files:
            logger.warning(f"No DOCX files found in '{self.data_folder}'")
            return documents
        
        logger.info(f"Found {len(docx_files)} DOCX files to process")
        
        for file_name in docx_files:
            file_path = os.path.join(self.data_folder, file_name)
            logger.info(f"Processing: {file_name}")
            
            text_content = self.extract_text_from_docx(file_path)
            
            if text_content:
                # Split the document into chunks
                chunks = self.text_splitter.split_text(text_content)
                
                for i, chunk in enumerate(chunks):
                    doc = LangChainDocument(
                        page_content=chunk,
                        metadata={
                            "source": file_name,
                            "chunk_id": i,
                            "file_path": file_path,
                            "total_chunks": len(chunks)
                        }
                    )
                    documents.append(doc)
                
                logger.info(f"Created {len(chunks)} chunks from {file_name}")
            else:
                logger.warning(f"No text content extracted from {file_name}")
        
        logger.info(f"Total documents loaded: {len(documents)}")
        return documents
    
    def get_file_info(self) -> Dict[str, int]:
        """Get information about files in the data folder."""
        info = {
            "total_files": 0,
            "docx_files": 0,
            "other_files": 0
        }
        
        if not os.path.exists(self.data_folder):
            return info
        
        files = os.listdir(self.data_folder)
        info["total_files"] = len(files)
        info["docx_files"] = len([f for f in files if f.endswith('.docx')])
        info["other_files"] = info["total_files"] - info["docx_files"]
        
        return info
    
    def load_single_document(self, file_path: str) -> List[LangChainDocument]:
        """Load a single DOCX document."""
        documents = []
        
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return documents
        
        if not file_path.endswith('.docx'):
            logger.error(f"File is not a DOCX file: {file_path}")
            return documents
        
        text_content = self.extract_text_from_docx(file_path)
        
        if text_content:
            chunks = self.text_splitter.split_text(text_content)
            
            for i, chunk in enumerate(chunks):
                doc = LangChainDocument(
                    page_content=chunk,
                    metadata={
                        "source": os.path.basename(file_path),
                        "chunk_id": i,
                        "file_path": file_path,
                        "total_chunks": len(chunks)
                    }
                )
                documents.append(doc)
        
        return documents